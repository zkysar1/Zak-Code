"""Create user-facing Word and Excel artifacts."""

from __future__ import annotations

import contextlib
import os
import re
import tempfile
from collections.abc import Mapping, Sequence
from pathlib import Path
from typing import Any

from zakcode.artifacts import ArtifactError, artifact_from_path
from zakcode.config import PermissionTier
from zakcode.tools.base import ConcurrencyClass, Tool, ToolContext, ToolResult, ToolSpec
from zakcode.tools.builtins._image_data import (
    format_from_suffix,
    sniff_image_format,
    supported_image_suffixes,
)
from zakcode.tools.builtins._safety import PathEscapeError, resolve_path
from zakcode.tools.builtins.images import render_chart_png

_EXCEL_FORMULA_PREFIXES = ("=", "+", "-", "@")
_INVALID_SHEET_CHARS = re.compile(r"[\[\]:*?/\\]")
_MAX_OFFICE_READ_BYTES = 50 * 1024 * 1024


def _with_suffix(path: str, suffix: str) -> str:
    candidate = Path(path)
    if candidate.suffix.lower() == suffix:
        return path
    if candidate.suffix:
        raise ValueError(f"path must end with {suffix}")
    return f"{path}{suffix}"


def _prepare_target(path: object, suffix: str, ctx: ToolContext) -> tuple[str, Path] | str:
    if not isinstance(path, str) or not path:
        return "'path' is required and must be a string."
    try:
        final_path = _with_suffix(path, suffix)
        resolved = resolve_path(final_path, ctx.workspace_root, ctx.extra_workspace_roots)
    except (PathEscapeError, ValueError) as exc:
        return str(exc)
    except Exception as exc:  # noqa: BLE001 - handlers must never raise
        return f"Failed to resolve path {path!r}: {exc}"
    if resolved.exists() and resolved.is_dir():
        return f"Path is a directory, not a file: {final_path}"
    return final_path, resolved


def _prepare_existing_file(
    path: object, suffixes: tuple[str, ...], ctx: ToolContext
) -> tuple[str, Path] | str:
    if not isinstance(path, str) or not path:
        return "'path' is required and must be a string."
    try:
        resolved = resolve_path(path, ctx.workspace_root, ctx.extra_workspace_roots)
    except PathEscapeError as exc:
        return str(exc)
    except Exception as exc:  # noqa: BLE001 - handlers must never raise
        return f"Failed to resolve path {path!r}: {exc}"
    if not resolved.is_file():
        return f"File not found: {path}"
    if resolved.suffix.lower() not in suffixes:
        return f"path must end with {' or '.join(suffixes)}"
    size = resolved.stat().st_size
    if size > _MAX_OFFICE_READ_BYTES:
        return f"file is too large ({size} bytes; max {_MAX_OFFICE_READ_BYTES})"
    return path, resolved


def _ensure_parent(resolved: Path) -> str | None:
    parent = resolved.parent
    if parent.exists() and not parent.is_dir():
        return f"Parent path is not a directory: {parent}"
    try:
        parent.mkdir(parents=True, exist_ok=True)
    except OSError as exc:
        return f"Could not create parent directory for {resolved}: {exc}"
    return None


def _replace_with_artifact(
    *,
    tmp_name: str,
    resolved: Path,
    ctx: ToolContext,
    tool_name: str,
    data: dict[str, Any],
    output: str,
) -> ToolResult:
    os.replace(tmp_name, resolved)
    artifacts = []
    with contextlib.suppress(ArtifactError, OSError):
        artifacts.append(
            artifact_from_path(
                resolved,
                workspace_root=ctx.workspace_root,
                created_by_tool=tool_name,
            )
        )
    if artifacts:
        data["artifact_id"] = artifacts[0].id
    return ToolResult.ok(output, data=data, artifacts=artifacts)


def _string_list(value: Any, *, field: str) -> list[str]:
    if not isinstance(value, list) or not all(isinstance(item, str) for item in value):
        raise ValueError(f"{field} must be a list of strings")
    return value


def _cell_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str | int | float | bool):
        return str(value)
    return str(value)


def _preview_text(value: Any) -> str:
    if value is None:
        return ""
    if hasattr(value, "isoformat"):
        return str(value.isoformat())
    return str(value).replace("\r", " ").replace("\n", " ")


def _bounded_int(value: Any, *, default: int, minimum: int, maximum: int, field: str) -> int:
    if value is None:
        return default
    if isinstance(value, bool) or not isinstance(value, int) or not minimum <= value <= maximum:
        raise ValueError(f"{field} must be an integer from {minimum} to {maximum}")
    return value


def _clip_lines(lines: list[str], max_chars: int) -> tuple[str, bool]:
    out: list[str] = []
    used = 0
    for line in lines:
        needed = len(line) + (1 if out else 0)
        if used + needed > max_chars:
            out.append("[truncated]")
            return "\n".join(out), True
        out.append(line)
        used += needed
    return "\n".join(out), False


def _excel_value(value: Any, *, allow_formulas: bool) -> str | int | float | bool | None:
    if value is None or isinstance(value, int | float | bool):
        return value
    if isinstance(value, str):
        if not allow_formulas and value.startswith(_EXCEL_FORMULA_PREFIXES):
            return "'" + value
        return value
    return str(value)


def _safe_sheet_name(raw: Any, fallback: str, used: set[str]) -> str:
    name = str(raw or fallback).strip()
    name = _INVALID_SHEET_CHARS.sub(" ", name)[:31].strip() or fallback
    base = name
    index = 2
    while name in used:
        suffix = f" {index}"
        name = f"{base[: 31 - len(suffix)]}{suffix}"
        index += 1
    used.add(name)
    return name


def _positive_int(value: Any, *, default: int, field: str) -> int:
    if value is None:
        return default
    if not isinstance(value, int) or value < 1:
        raise ValueError(f"{field} must be a positive integer")
    return value


class ReadDocxTool(Tool):
    """Read a Microsoft Word document into capped text and table previews."""

    spec = ToolSpec(
        name="read_docx",
        description=(
            "Read a Microsoft Word .docx file in the workspace. Extracts capped paragraph "
            "text, table previews, and image counts for analysis without dumping binary data."
        ),
        parameters={
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "Workspace .docx path to read."},
                "max_paragraphs": {"type": "integer", "minimum": 1, "maximum": 200},
                "max_tables": {"type": "integer", "minimum": 0, "maximum": 20},
                "max_table_rows": {"type": "integer", "minimum": 1, "maximum": 50},
                "max_chars": {"type": "integer", "minimum": 1000, "maximum": 50000},
            },
            "required": ["path"],
        },
        required_permission=PermissionTier.READ_ONLY,
        concurrency=ConcurrencyClass.READ_ONLY_SAFE,
    )

    async def execute(self, args: dict[str, Any], ctx: ToolContext) -> ToolResult:
        target = _prepare_existing_file(args.get("path"), (".docx",), ctx)
        if isinstance(target, str):
            return ToolResult.error(target)
        display_path, resolved = target
        try:
            max_paragraphs = _bounded_int(
                args.get("max_paragraphs"),
                default=80,
                minimum=1,
                maximum=200,
                field="max_paragraphs",
            )
            max_tables = _bounded_int(
                args.get("max_tables"), default=5, minimum=0, maximum=20, field="max_tables"
            )
            max_table_rows = _bounded_int(
                args.get("max_table_rows"),
                default=10,
                minimum=1,
                maximum=50,
                field="max_table_rows",
            )
            max_chars = _bounded_int(
                args.get("max_chars"),
                default=12000,
                minimum=1000,
                maximum=50000,
                field="max_chars",
            )
        except ValueError as exc:
            return ToolResult.error(str(exc))

        try:
            from docx import Document
        except ImportError:
            return ToolResult.error("python-docx is required for read_docx; run `uv sync`.")

        try:
            doc = Document(str(resolved))
            paragraphs = [
                paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()
            ]
            shown_paragraphs = paragraphs[:max_paragraphs]
            table_summaries: list[dict[str, Any]] = []
            lines = [
                f"Word document: {display_path}",
                f"Paragraphs: {len(paragraphs)} ({len(shown_paragraphs)} shown)",
                f"Tables: {len(doc.tables)}",
                f"Images: {len(doc.inline_shapes)}",
            ]
            if shown_paragraphs:
                lines.extend(["", "Text:", *shown_paragraphs])

            for table_index, table in enumerate(doc.tables[:max_tables], start=1):
                rows: list[list[str]] = []
                for row in table.rows[:max_table_rows]:
                    rows.append([_preview_text(cell.text) for cell in row.cells])
                summary = {
                    "index": table_index,
                    "row_count": len(table.rows),
                    "column_count": len(table.columns),
                    "rows": rows,
                }
                table_summaries.append(summary)
                lines.extend(
                    [
                        "",
                        (
                            f"Table {table_index}: {summary['row_count']} rows x "
                            f"{summary['column_count']} columns ({len(rows)} rows shown)"
                        ),
                    ]
                )
                lines.extend(" | ".join(row) for row in rows)

            output, truncated = _clip_lines(lines, max_chars)
            data = {
                "path": str(resolved),
                "bytes": resolved.stat().st_size,
                "paragraph_count": len(paragraphs),
                "paragraphs": shown_paragraphs,
                "table_count": len(doc.tables),
                "tables": table_summaries,
                "image_count": len(doc.inline_shapes),
                "truncated": truncated,
            }
            return ToolResult.ok(output, data=data)
        except Exception as exc:  # noqa: BLE001 - handlers must never raise
            return ToolResult.error(f"Failed to read Word document {display_path!r}: {exc}")


class ReadXlsxTool(Tool):
    """Read a Microsoft Excel workbook into capped sheet previews."""

    spec = ToolSpec(
        name="read_xlsx",
        description=(
            "Read a Microsoft Excel .xlsx or .xlsm workbook in the workspace. Returns sheet "
            "dimensions, table/chart counts, merged-cell counts, and capped cell previews."
        ),
        parameters={
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "Workspace .xlsx/.xlsm path to read."},
                "max_sheets": {"type": "integer", "minimum": 1, "maximum": 50},
                "max_rows": {"type": "integer", "minimum": 1, "maximum": 200},
                "max_columns": {"type": "integer", "minimum": 1, "maximum": 50},
                "max_chars": {"type": "integer", "minimum": 1000, "maximum": 50000},
            },
            "required": ["path"],
        },
        required_permission=PermissionTier.READ_ONLY,
        concurrency=ConcurrencyClass.READ_ONLY_SAFE,
    )

    async def execute(self, args: dict[str, Any], ctx: ToolContext) -> ToolResult:
        target = _prepare_existing_file(args.get("path"), (".xlsx", ".xlsm"), ctx)
        if isinstance(target, str):
            return ToolResult.error(target)
        display_path, resolved = target
        try:
            max_sheets = _bounded_int(
                args.get("max_sheets"), default=10, minimum=1, maximum=50, field="max_sheets"
            )
            max_rows = _bounded_int(
                args.get("max_rows"), default=20, minimum=1, maximum=200, field="max_rows"
            )
            max_columns = _bounded_int(
                args.get("max_columns"),
                default=12,
                minimum=1,
                maximum=50,
                field="max_columns",
            )
            max_chars = _bounded_int(
                args.get("max_chars"),
                default=20000,
                minimum=1000,
                maximum=50000,
                field="max_chars",
            )
        except ValueError as exc:
            return ToolResult.error(str(exc))

        try:
            from openpyxl import load_workbook
        except ImportError:
            return ToolResult.error("openpyxl is required for read_xlsx; run `uv sync`.")

        workbook = None
        try:
            workbook = load_workbook(
                resolved,
                data_only=False,
                read_only=False,
                keep_vba=resolved.suffix.lower() == ".xlsm",
                keep_links=False,
            )
            shown_sheets = workbook.worksheets[:max_sheets]
            lines = [
                f"Excel workbook: {display_path}",
                f"Sheets: {len(workbook.worksheets)} ({len(shown_sheets)} shown)",
            ]
            sheet_summaries: list[dict[str, Any]] = []
            for sheet in shown_sheets:
                row_limit = min(max_rows, sheet.max_row or 0)
                column_limit = min(max_columns, sheet.max_column or 0)
                rows: list[list[str]] = []
                formula_count = 0
                for row in sheet.iter_rows(
                    min_row=1,
                    max_row=row_limit,
                    min_col=1,
                    max_col=column_limit,
                ):
                    preview_row: list[str] = []
                    for cell in row:
                        value = cell.value
                        if isinstance(value, str) and value.startswith("="):
                            formula_count += 1
                        preview_row.append(_preview_text(value))
                    rows.append(preview_row)

                summary = {
                    "name": sheet.title,
                    "row_count": sheet.max_row,
                    "column_count": sheet.max_column,
                    "rows": rows,
                    "tables": len(sheet.tables),
                    "charts": len(sheet._charts),
                    "merged_cells": len(sheet.merged_cells.ranges),
                    "formula_cells_in_preview": formula_count,
                }
                sheet_summaries.append(summary)
                lines.extend(
                    [
                        "",
                        (
                            f"Sheet {sheet.title!r}: {sheet.max_row} rows x {sheet.max_column} "
                            f"columns; tables={summary['tables']}; charts={summary['charts']}; "
                            f"merged={summary['merged_cells']}; formulas in preview={formula_count}"
                        ),
                        f"Preview: {len(rows)} rows x {column_limit} columns",
                    ]
                )
                lines.extend(" | ".join(row) for row in rows)

            output, truncated = _clip_lines(lines, max_chars)
            data = {
                "path": str(resolved),
                "bytes": resolved.stat().st_size,
                "sheet_count": len(workbook.worksheets),
                "sheets": sheet_summaries,
                "truncated": truncated,
            }
            return ToolResult.ok(output, data=data)
        except Exception as exc:  # noqa: BLE001 - handlers must never raise
            return ToolResult.error(f"Failed to read Excel workbook {display_path!r}: {exc}")
        finally:
            if workbook is not None:
                with contextlib.suppress(Exception):
                    workbook.close()


class CreateDocxTool(Tool):
    """Create a Microsoft Word document as a downloadable artifact."""

    spec = ToolSpec(
        name="create_docx",
        description=(
            "Create a Microsoft Word .docx file within the workspace. Supports title, "
            "headings, paragraphs, bullet/numbered lists, tables, images, chart images, "
            "and page breaks. Returns a downloadable artifact."
        ),
        parameters={
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "Destination path. .docx is appended when no suffix is given.",
                },
                "title": {"type": "string", "description": "Optional document title."},
                "blocks": {
                    "type": "array",
                    "description": "Ordered Word content blocks.",
                    "items": {
                        "type": "object",
                        "properties": {
                            "type": {
                                "type": "string",
                                "enum": [
                                    "heading",
                                    "paragraph",
                                    "bullet_list",
                                    "numbered_list",
                                    "table",
                                    "image",
                                    "chart",
                                    "page_break",
                                ],
                            },
                            "text": {"type": "string"},
                            "level": {"type": "integer", "minimum": 1, "maximum": 4},
                            "items": {"type": "array", "items": {"type": "string"}},
                            "headers": {"type": "array", "items": {"type": "string"}},
                            "rows": {
                                "type": "array",
                                "items": {"type": "array", "items": {}},
                            },
                            "path": {
                                "type": "string",
                                "description": "Workspace path for an image block.",
                            },
                            "width_inches": {
                                "type": "number",
                                "minimum": 0.1,
                                "maximum": 10,
                            },
                            "caption": {"type": "string"},
                            "chart_type": {
                                "type": "string",
                                "enum": ["bar", "line", "pie"],
                            },
                            "title": {"type": "string"},
                            "data": {
                                "type": "array",
                                "maxItems": 100,
                                "items": {
                                    "type": "object",
                                    "properties": {
                                        "label": {"type": "string"},
                                        "value": {"type": "number", "minimum": 0},
                                    },
                                    "required": ["label", "value"],
                                },
                            },
                            "width": {"type": "integer", "minimum": 320, "maximum": 2400},
                            "height": {"type": "integer", "minimum": 240, "maximum": 1600},
                        },
                        "required": ["type"],
                    },
                },
            },
            "required": ["path"],
        },
        required_permission=PermissionTier.WORKSPACE_WRITE,
        concurrency=ConcurrencyClass.PATH_SCOPED,
    )

    async def execute(self, args: dict, ctx: ToolContext) -> ToolResult:
        target = _prepare_target(args.get("path"), ".docx", ctx)
        if isinstance(target, str):
            return ToolResult.error(target)
        final_path, resolved = target
        if (parent_error := _ensure_parent(resolved)) is not None:
            return ToolResult.error(parent_error)

        try:
            from docx import Document
        except ImportError:
            return ToolResult.error("python-docx is required for create_docx; run `uv sync`.")

        title = args.get("title")
        blocks = args.get("blocks", [])
        if title is not None and not isinstance(title, str):
            return ToolResult.error("'title' must be a string when provided.")
        if not isinstance(blocks, list):
            return ToolResult.error("'blocks' must be a list when provided.")
        if not title and not blocks:
            return ToolResult.error("Provide 'title' or at least one content block.")

        try:
            doc = Document()
            if title:
                doc.add_heading(title, level=0)
            self._add_blocks(doc, blocks, ctx)

            fd, tmp_name = tempfile.mkstemp(
                prefix=f".{resolved.name}.", suffix=".tmp", dir=str(resolved.parent)
            )
            os.close(fd)
            try:
                doc.save(tmp_name)
                return _replace_with_artifact(
                    tmp_name=tmp_name,
                    resolved=resolved,
                    ctx=ctx,
                    tool_name=self.spec.name,
                    data={"path": str(resolved), "bytes": Path(tmp_name).stat().st_size},
                    output=f"Created Word document {final_path}",
                )
            except Exception:
                with contextlib.suppress(OSError):
                    os.unlink(tmp_name)
                raise
        except Exception as exc:  # noqa: BLE001 - handlers must never raise
            return ToolResult.error(f"Failed to create Word document {final_path!r}: {exc}")

    def _add_blocks(self, doc: Any, blocks: list[Any], ctx: ToolContext) -> None:
        for index, block in enumerate(blocks, start=1):
            if not isinstance(block, Mapping):
                raise ValueError(f"blocks[{index}] must be an object")
            kind = block.get("type")
            if kind == "heading":
                text = block.get("text")
                if not isinstance(text, str):
                    raise ValueError(f"blocks[{index}].text must be a string")
                level = block.get("level", 1)
                if not isinstance(level, int) or not 1 <= level <= 4:
                    raise ValueError(f"blocks[{index}].level must be an integer from 1 to 4")
                doc.add_heading(text, level=level)
            elif kind == "paragraph":
                text = block.get("text")
                if not isinstance(text, str):
                    raise ValueError(f"blocks[{index}].text must be a string")
                doc.add_paragraph(text)
            elif kind in {"bullet_list", "numbered_list"}:
                items = _string_list(block.get("items"), field=f"blocks[{index}].items")
                style = "List Bullet" if kind == "bullet_list" else "List Number"
                for item in items:
                    doc.add_paragraph(item, style=style)
            elif kind == "table":
                self._add_table(doc, block, index)
            elif kind == "image":
                self._add_image(doc, block, index, ctx)
            elif kind == "chart":
                self._add_chart(doc, block, index, ctx)
            elif kind == "page_break":
                doc.add_page_break()
            else:
                raise ValueError(f"blocks[{index}].type is not supported: {kind!r}")

    def _add_table(self, doc: Any, block: Mapping[str, Any], index: int) -> None:
        headers_raw = block.get("headers", [])
        headers = _string_list(headers_raw, field=f"blocks[{index}].headers") if headers_raw else []
        rows_raw = block.get("rows")
        if not isinstance(rows_raw, list) or not all(isinstance(row, list) for row in rows_raw):
            raise ValueError(f"blocks[{index}].rows must be a list of row arrays")
        rows: list[list[Any]] = rows_raw
        column_count = max([len(headers), *(len(row) for row in rows)], default=0)
        if column_count == 0:
            raise ValueError(f"blocks[{index}] table must have headers or rows")

        table = doc.add_table(rows=len(rows) + (1 if headers else 0), cols=column_count)
        table.style = "Table Grid"
        offset = 0
        if headers:
            for col, value in enumerate(headers):
                table.cell(0, col).text = value
            offset = 1
        for row_idx, row in enumerate(rows, start=offset):
            for col, value in enumerate(row):
                table.cell(row_idx, col).text = _cell_text(value)

    def _add_image(self, doc: Any, block: Mapping[str, Any], index: int, ctx: ToolContext) -> None:
        from docx.shared import Inches

        source = block.get("path")
        if not isinstance(source, str) or not source:
            raise ValueError(f"blocks[{index}].path must be a string for image blocks")
        image_path = resolve_path(source, ctx.workspace_root, ctx.extra_workspace_roots)
        if not image_path.is_file():
            raise ValueError(f"blocks[{index}].path is not a file: {source}")

        suffix_format = format_from_suffix(image_path)
        if suffix_format is None:
            raise ValueError(
                f"blocks[{index}].path must end with one of {', '.join(supported_image_suffixes())}"
            )
        with image_path.open("rb") as handle:
            detected = sniff_image_format(handle.read(16))
        if detected is None or detected.mime_type != suffix_format.mime_type:
            raise ValueError(f"blocks[{index}].path is not a valid {suffix_format.mime_type} image")

        width = block.get("width_inches", 5.5)
        if (
            not isinstance(width, int | float)
            or isinstance(width, bool)
            or width <= 0
            or width > 10
        ):
            raise ValueError(f"blocks[{index}].width_inches must be a number from 0.1 to 10")
        doc.add_picture(str(image_path), width=Inches(float(width)))

        caption = block.get("caption")
        if caption is not None:
            if not isinstance(caption, str):
                raise ValueError(f"blocks[{index}].caption must be a string")
            if caption:
                doc.add_paragraph(caption)

    def _add_chart(self, doc: Any, block: Mapping[str, Any], index: int, ctx: ToolContext) -> None:
        from docx.shared import Inches

        width = block.get("width_inches", 5.5)
        if (
            not isinstance(width, int | float)
            or isinstance(width, bool)
            or width <= 0
            or width > 10
        ):
            raise ValueError(f"blocks[{index}].width_inches must be a number from 0.1 to 10")

        try:
            chart_bytes = render_chart_png(block)
        except ValueError as exc:
            raise ValueError(f"blocks[{index}] chart is invalid: {exc}") from exc

        fd, tmp_name = tempfile.mkstemp(
            prefix=".zakcode-chart-", suffix=".png", dir=str(ctx.workspace_root)
        )
        try:
            with os.fdopen(fd, "wb") as handle:
                handle.write(chart_bytes)
            doc.add_picture(tmp_name, width=Inches(float(width)))
        finally:
            with contextlib.suppress(OSError):
                os.unlink(tmp_name)

        caption = block.get("caption")
        if caption is not None:
            if not isinstance(caption, str):
                raise ValueError(f"blocks[{index}].caption must be a string")
            if caption:
                doc.add_paragraph(caption)


class CreateXlsxTool(Tool):
    """Create a Microsoft Excel workbook as a downloadable artifact."""

    spec = ToolSpec(
        name="create_xlsx",
        description=(
            "Create a Microsoft Excel .xlsx workbook within the workspace. Supports multiple "
            "sheets, header formatting, optional Excel tables, and simple bar/line/pie charts. "
            "Returns a downloadable artifact."
        ),
        parameters={
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "Destination path. .xlsx is appended when no suffix is given.",
                },
                "sheets": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "name": {"type": "string"},
                            "headers": {"type": "array", "items": {"type": "string"}},
                            "rows": {"type": "array", "items": {}},
                            "table": {"type": "boolean"},
                            "freeze_header": {"type": "boolean"},
                            "allow_formulas": {"type": "boolean"},
                            "charts": {
                                "type": "array",
                                "items": {
                                    "type": "object",
                                    "properties": {
                                        "type": {
                                            "type": "string",
                                            "enum": ["bar", "line", "pie"],
                                        },
                                        "title": {"type": "string"},
                                        "categories_col": {"type": "integer", "minimum": 1},
                                        "values_col": {"type": "integer", "minimum": 1},
                                        "start_row": {"type": "integer", "minimum": 1},
                                        "end_row": {"type": "integer", "minimum": 1},
                                        "position": {"type": "string"},
                                    },
                                },
                            },
                        },
                        "required": ["rows"],
                    },
                },
            },
            "required": ["path", "sheets"],
        },
        required_permission=PermissionTier.WORKSPACE_WRITE,
        concurrency=ConcurrencyClass.PATH_SCOPED,
    )

    async def execute(self, args: dict, ctx: ToolContext) -> ToolResult:
        target = _prepare_target(args.get("path"), ".xlsx", ctx)
        if isinstance(target, str):
            return ToolResult.error(target)
        final_path, resolved = target
        if (parent_error := _ensure_parent(resolved)) is not None:
            return ToolResult.error(parent_error)

        try:
            from openpyxl import Workbook
        except ImportError:
            return ToolResult.error("openpyxl is required for create_xlsx; run `uv sync`.")

        sheets = args.get("sheets")
        if not isinstance(sheets, list) or not sheets:
            return ToolResult.error("'sheets' is required and must be a non-empty list.")

        try:
            wb = Workbook()
            used_names: set[str] = set()
            for index, sheet in enumerate(sheets, start=1):
                if not isinstance(sheet, Mapping):
                    raise ValueError(f"sheets[{index}] must be an object")
                ws = wb.active if index == 1 else wb.create_sheet()
                ws.title = _safe_sheet_name(sheet.get("name"), f"Sheet{index}", used_names)
                self._populate_sheet(ws, sheet, index)

            fd, tmp_name = tempfile.mkstemp(
                prefix=f".{resolved.name}.", suffix=".tmp", dir=str(resolved.parent)
            )
            os.close(fd)
            try:
                wb.save(tmp_name)
                return _replace_with_artifact(
                    tmp_name=tmp_name,
                    resolved=resolved,
                    ctx=ctx,
                    tool_name=self.spec.name,
                    data={
                        "path": str(resolved),
                        "bytes": Path(tmp_name).stat().st_size,
                        "sheets": [ws.title for ws in wb.worksheets],
                    },
                    output=f"Created Excel workbook {final_path}",
                )
            except Exception:
                with contextlib.suppress(OSError):
                    os.unlink(tmp_name)
                raise
        except Exception as exc:  # noqa: BLE001 - handlers must never raise
            return ToolResult.error(f"Failed to create Excel workbook {final_path!r}: {exc}")

    def _populate_sheet(self, ws: Any, sheet: Mapping[str, Any], index: int) -> None:
        from openpyxl.styles import Font
        from openpyxl.utils import get_column_letter
        from openpyxl.worksheet.table import Table, TableStyleInfo

        allow_formulas = bool(sheet.get("allow_formulas", False))
        headers, rows = self._normalize_rows(sheet, index, allow_formulas=allow_formulas)
        if headers:
            ws.append(headers)
            for cell in ws[1]:
                cell.font = Font(bold=True)
        for row in rows:
            ws.append(row)

        if headers and sheet.get("freeze_header", True):
            ws.freeze_panes = "A2"
        for col_idx, width in enumerate(self._column_widths(headers, rows), start=1):
            ws.column_dimensions[get_column_letter(col_idx)].width = min(max(width, 10), 42)
        if headers and rows and sheet.get("table", False):
            end_col = get_column_letter(len(headers))
            ref = f"A1:{end_col}{len(rows) + 1}"
            table = Table(displayName=f"Table{index}", ref=ref)
            table.tableStyleInfo = TableStyleInfo(
                name="TableStyleMedium2",
                showFirstColumn=False,
                showLastColumn=False,
                showRowStripes=True,
                showColumnStripes=False,
            )
            ws.add_table(table)
        self._add_charts(ws, sheet.get("charts", []), has_headers=bool(headers))

    def _normalize_rows(
        self, sheet: Mapping[str, Any], index: int, *, allow_formulas: bool
    ) -> tuple[list[str], list[list[Any]]]:
        headers_raw = sheet.get("headers")
        rows_raw = sheet.get("rows")
        if headers_raw is not None:
            headers = _string_list(headers_raw, field=f"sheets[{index}].headers")
        else:
            headers = []
        if not isinstance(rows_raw, list):
            raise ValueError(f"sheets[{index}].rows must be a list")

        if any(isinstance(row, Mapping) for row in rows_raw):
            if not all(isinstance(row, Mapping) for row in rows_raw):
                raise ValueError(f"sheets[{index}].rows must not mix objects and arrays")
            if not headers:
                seen: dict[str, None] = {}
                for row in rows_raw:
                    for key in row:
                        seen.setdefault(str(key), None)
                headers = list(seen)
            object_rows = [
                [_excel_value(row.get(header), allow_formulas=allow_formulas) for header in headers]
                for row in rows_raw
            ]
            return headers, object_rows

        array_rows: list[list[Any]] = []
        for row_index, row in enumerate(rows_raw, start=1):
            if not isinstance(row, Sequence) or isinstance(row, str | bytes):
                raise ValueError(f"sheets[{index}].rows[{row_index}] must be an array")
            array_rows.append([_excel_value(value, allow_formulas=allow_formulas) for value in row])
        return headers, array_rows

    def _column_widths(self, headers: list[str], rows: list[list[Any]]) -> list[int]:
        column_count = max([len(headers), *(len(row) for row in rows)], default=1)
        widths: list[int] = []
        for col_idx in range(column_count):
            values = [headers[col_idx]] if col_idx < len(headers) else []
            values.extend(
                str(row[col_idx]) for row in rows if col_idx < len(row) and row[col_idx] is not None
            )
            widths.append(max([len(value) for value in values], default=0) + 2)
        return widths

    def _add_charts(self, ws: Any, charts: Any, *, has_headers: bool) -> None:
        if charts in (None, []):
            return
        if not isinstance(charts, list):
            raise ValueError("charts must be a list when provided")
        for index, chart_config in enumerate(charts, start=1):
            if not isinstance(chart_config, Mapping):
                raise ValueError(f"charts[{index}] must be an object")
            self._add_chart(ws, chart_config, index=index, has_headers=has_headers)

    def _add_chart(
        self, ws: Any, chart_config: Mapping[str, Any], *, index: int, has_headers: bool
    ) -> None:
        from openpyxl.chart import BarChart, LineChart, PieChart, Reference
        from openpyxl.utils import get_column_letter

        chart_type = str(chart_config.get("type", "bar")).lower()
        if chart_type == "bar":
            chart = BarChart()
        elif chart_type == "line":
            chart = LineChart()
        elif chart_type == "pie":
            chart = PieChart()
        else:
            raise ValueError(f"charts[{index}].type must be bar, line, or pie")

        categories_col = _positive_int(
            chart_config.get("categories_col"), default=1, field=f"charts[{index}].categories_col"
        )
        values_col = _positive_int(
            chart_config.get("values_col"), default=2, field=f"charts[{index}].values_col"
        )
        data_start = 2 if has_headers else 1
        start_row = _positive_int(
            chart_config.get("start_row"), default=data_start, field=f"charts[{index}].start_row"
        )
        end_row = _positive_int(
            chart_config.get("end_row"), default=ws.max_row, field=f"charts[{index}].end_row"
        )
        if end_row < start_row:
            raise ValueError(f"charts[{index}].end_row must be >= start_row")
        if max(categories_col, values_col) > ws.max_column or end_row > ws.max_row:
            raise ValueError(f"charts[{index}] references cells outside the sheet data")

        title = chart_config.get("title")
        if title is not None:
            if not isinstance(title, str):
                raise ValueError(f"charts[{index}].title must be a string")
            chart.title = title

        data_min_row = 1 if has_headers and start_row > 1 else start_row
        data = Reference(ws, min_col=values_col, min_row=data_min_row, max_row=end_row)
        categories = Reference(ws, min_col=categories_col, min_row=start_row, max_row=end_row)
        chart.add_data(data, titles_from_data=data_min_row < start_row)
        chart.set_categories(categories)

        position = chart_config.get("position")
        if position is not None and not isinstance(position, str):
            raise ValueError(f"charts[{index}].position must be a cell reference string")
        anchor = position or f"{get_column_letter(ws.max_column + 2)}2"
        ws.add_chart(chart, anchor)
