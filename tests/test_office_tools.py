"""Tests for Word/Excel artifact producer tools."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from openpyxl import load_workbook

from zakcode.tools.base import ToolContext
from zakcode.tools.builtins.office import CreateDocxTool, CreateXlsxTool, ReadDocxTool, ReadXlsxTool


async def test_create_docx_generates_downloadable_word_document(tmp_path: Path) -> None:
    ctx = ToolContext(workspace_root=tmp_path)
    result = await CreateDocxTool().execute(
        {
            "path": "reports/family-plan",
            "title": "Family Plan",
            "blocks": [
                {"type": "heading", "text": "Appointments", "level": 1},
                {"type": "paragraph", "text": "Bring the insurance card."},
                {"type": "bullet_list", "items": ["Call the clinic", "Pack documents"]},
                {
                    "type": "table",
                    "headers": ["Task", "Owner"],
                    "rows": [["Call", "Sam"], ["Drive", "Lee"]],
                },
            ],
        },
        ctx,
    )

    assert not result.is_error, result.output
    path = tmp_path / "reports" / "family-plan.docx"
    assert path.exists()
    assert len(result.artifacts) == 1
    assert result.artifacts[0].path == "reports/family-plan.docx"
    assert result.artifacts[0].kind == "document"
    assert result.artifacts[0].created_by_tool == "create_docx"
    assert result.data is not None
    assert result.data["artifact_id"] == result.artifacts[0].id

    doc = Document(path)
    paragraphs = [paragraph.text for paragraph in doc.paragraphs]
    assert "Family Plan" in paragraphs
    assert "Appointments" in paragraphs
    assert "Bring the insurance card." in paragraphs
    assert "Call the clinic" in paragraphs
    assert doc.tables[0].cell(0, 0).text == "Task"
    assert doc.tables[0].cell(2, 1).text == "Lee"


async def test_create_docx_embeds_generated_chart_image(tmp_path: Path) -> None:
    ctx = ToolContext(workspace_root=tmp_path)
    result = await CreateDocxTool().execute(
        {
            "path": "reports/with-chart.docx",
            "title": "Chart Report",
            "blocks": [
                {
                    "type": "chart",
                    "chart_type": "line",
                    "title": "Quarterly Revenue",
                    "width": 640,
                    "height": 360,
                    "width_inches": 5,
                    "caption": "Quarterly revenue trend.",
                    "data": [
                        {"label": "Q1", "value": 100},
                        {"label": "Q2", "value": 130},
                        {"label": "Q3", "value": 125},
                        {"label": "Q4", "value": 170},
                    ],
                }
            ],
        },
        ctx,
    )

    assert not result.is_error, result.output
    doc = Document(tmp_path / "reports" / "with-chart.docx")
    assert len(doc.inline_shapes) == 1
    assert "Quarterly revenue trend." in [paragraph.text for paragraph in doc.paragraphs]


async def test_read_docx_extracts_text_tables_and_image_count(tmp_path: Path) -> None:
    ctx = ToolContext(workspace_root=tmp_path)
    created = await CreateDocxTool().execute(
        {
            "path": "reports/read-me.docx",
            "title": "Read Me",
            "blocks": [
                {"type": "paragraph", "text": "This is the key sentence."},
                {"type": "table", "headers": ["Name", "Value"], "rows": [["A", 1], ["B", 2]]},
            ],
        },
        ctx,
    )
    assert not created.is_error, created.output

    result = await ReadDocxTool().execute({"path": "reports/read-me.docx"}, ctx)

    assert not result.is_error, result.output
    assert "This is the key sentence." in result.output
    assert "Table 1: 3 rows x 2 columns" in result.output
    assert result.data is not None
    assert result.data["paragraph_count"] >= 2
    assert result.data["table_count"] == 1
    assert result.data["image_count"] == 0
    assert result.data["tables"][0]["rows"][0] == ["Name", "Value"]


async def test_office_readers_reject_path_escape(tmp_path: Path) -> None:
    ctx = ToolContext(workspace_root=tmp_path)

    docx_result = await ReadDocxTool().execute({"path": "../outside.docx"}, ctx)
    xlsx_result = await ReadXlsxTool().execute({"path": "../outside.xlsx"}, ctx)

    assert docx_result.is_error
    assert xlsx_result.is_error
    assert "outside" in docx_result.output.lower()
    assert "outside" in xlsx_result.output.lower()


async def test_create_xlsx_generates_workbook_table_and_chart(tmp_path: Path) -> None:
    ctx = ToolContext(workspace_root=tmp_path)
    result = await CreateXlsxTool().execute(
        {
            "path": "reports/budget",
            "sheets": [
                {
                    "name": "Budget",
                    "headers": ["Month", "Amount"],
                    "rows": [["Jan", 120], ["Feb", 160], ["Mar", 140]],
                    "table": True,
                    "charts": [
                        {
                            "type": "bar",
                            "title": "Monthly Amount",
                            "categories_col": 1,
                            "values_col": 2,
                            "position": "D2",
                        }
                    ],
                },
                {
                    "name": "Notes",
                    "rows": [{"Topic": "Next step", "Value": "Review with Grandma"}],
                },
            ],
        },
        ctx,
    )

    assert not result.is_error, result.output
    path = tmp_path / "reports" / "budget.xlsx"
    assert path.exists()
    assert len(result.artifacts) == 1
    assert result.artifacts[0].path == "reports/budget.xlsx"
    assert result.artifacts[0].kind == "spreadsheet"
    assert result.artifacts[0].created_by_tool == "create_xlsx"
    assert result.data is not None
    assert result.data["sheets"] == ["Budget", "Notes"]
    assert result.data["artifact_id"] == result.artifacts[0].id

    workbook = load_workbook(path)
    sheet = workbook["Budget"]
    assert sheet["A1"].value == "Month"
    assert sheet["B3"].value == 160
    assert sheet.freeze_panes == "A2"
    assert "Table1" in sheet.tables
    assert len(sheet._charts) == 1
    assert workbook["Notes"]["A1"].value == "Topic"
    assert workbook["Notes"]["B2"].value == "Review with Grandma"


async def test_read_xlsx_extracts_sheet_preview_and_chart_count(tmp_path: Path) -> None:
    ctx = ToolContext(workspace_root=tmp_path)
    created = await CreateXlsxTool().execute(
        {
            "path": "reports/read-budget.xlsx",
            "sheets": [
                {
                    "name": "Budget",
                    "headers": ["Month", "Amount"],
                    "rows": [["Jan", 120], ["Feb", 160]],
                    "table": True,
                    "charts": [{"type": "line", "categories_col": 1, "values_col": 2}],
                }
            ],
        },
        ctx,
    )
    assert not created.is_error, created.output

    result = await ReadXlsxTool().execute({"path": "reports/read-budget.xlsx"}, ctx)

    assert not result.is_error, result.output
    assert "Sheet 'Budget': 3 rows x 2 columns" in result.output
    assert "Jan | 120" in result.output
    assert result.data is not None
    assert result.data["sheet_count"] == 1
    assert result.data["sheets"][0]["tables"] == 1
    assert result.data["sheets"][0]["charts"] == 1
    assert result.data["sheets"][0]["rows"][0] == ["Month", "Amount"]


async def test_create_xlsx_escapes_formulas_by_default(tmp_path: Path) -> None:
    ctx = ToolContext(workspace_root=tmp_path)
    result = await CreateXlsxTool().execute(
        {
            "path": "formula-safe.xlsx",
            "sheets": [{"headers": ["Label"], "rows": [['=HYPERLINK("http://x")']]}],
        },
        ctx,
    )

    assert not result.is_error, result.output
    value = load_workbook(tmp_path / "formula-safe.xlsx").active["A2"].value
    assert value == '\'=HYPERLINK("http://x")'
